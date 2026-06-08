# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 设置标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '黑体'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)

def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
    return h

def add_p(doc, text, bold=False, size=12, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    return p

def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_screenshot_placeholder(doc, text):
    """添加截图占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('【' + text + '】')
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return p

# ==================== 封面 ====================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('多跳推理问答数据检索与聚类分析')
run.font.name = '黑体'
run.font.size = Pt(26)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 基于 HotpotQA 数据集 ——')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('学院：计算机与人工智能学院\n专业班级：数据科学与大数据技术 2302 班\n姓名：葛超瑜\n学号：2023XXXXXXXX\n完成日期：2026年6月8日')
run.font.name = '宋体'
run.font.size = Pt(14)

doc.add_page_break()

# ==================== 目录 ====================
add_h(doc, '目  录', 1)
toc_items = [
    ('一、实验背景与目的', 3),
    ('二、数据集介绍', 4),
    ('  2.1 HotpotQA 数据集概述', 4),
    ('  2.2 数据集结构', 5),
    ('三、数据获取与预处理', 6),
    ('  3.1 数据获取过程', 6),
    ('  3.2 数据预处理', 7),
    ('四、多跳检索系统设计', 9),
    ('  4.1 系统架构', 9),
    ('  4.2 多跳检索功能', 10),
    ('  4.3 检索示例分析', 11),
    ('五、聚类分析方法', 12),
    ('  5.1 TF-IDF 特征提取', 12),
    ('  5.2 K-Means 聚类', 13),
    ('  5.3 PCA 降维可视化', 14),
    ('六、可视化设计', 15),
    ('七、Web 应用部署', 16),
    ('八、实验结果与分析', 17),
    ('  8.1 聚类效果分析', 17),
    ('  8.2 多跳检索效果分析', 18),
    ('  8.3 数据集统计分析', 18),
    ('九、实验总结', 19),
    ('附录：截图清单', 20),
]
for item_text, page_num in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item_text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    # 添加制表符和页码
    tab_run = p.add_run('  ' + '.' * (50 - len(item_text)) + '  ' + str(page_num))
    tab_run.font.name = '宋体'
    tab_run.font.size = Pt(10)
    tab_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ==================== 一、实验背景与目的 ====================
add_h(doc, '一、实验背景与目的', 1)

add_p(doc, '随着人工智能与自然语言处理技术的快速发展，问答系统（Question Answering, QA）已成为该领域的核心研究方向之一。传统的单跳问答（Single-hop QA）只需要从单篇文档中找到答案，而多跳推理问答（Multi-hop QA）则要求系统能够从多篇文档中检索相关信息，并进行跨文档的逻辑推理，才能得到最终答案。这种多跳推理能力更接近人类的阅读推理过程，具有重要的研究价值和应用前景。', indent=True)

add_p(doc, 'HotpotQA 是由卡内基梅隆大学和斯坦福大学的研究团队于 2018 年发布的多跳推理问答数据集，在自然语言处理领域具有广泛的影响力。该数据集精心设计了需要跨越多个文档进行推理的问题，并提供句子级别的支持事实标注，为研究多跳推理机制提供了重要的基准测试平台。', indent=True)

add_p(doc, '本实验旨在通过实际动手操作，完成以下目标：')
objectives = [
    '了解多跳推理问答数据集（HotpotQA）的数据结构、特点和标注方式',
    '掌握使用 Python 进行数据获取、解析和预处理的基本方法',
    '学习 TF-IDF 文本特征提取技术的原理和应用',
    '掌握 K-Means 聚类算法的基本原理及其在文本分析中的应用',
    '学习 PCA 降维方法及其在数据可视化中的作用',
    '构建一个支持多跳检索、聚类分析和可视化的 Web 交互系统',
    '掌握使用 GitHub Pages 部署静态网站的方法',
]
for obj in objectives:
    add_bullet(doc, '  ● ' + obj)

doc.add_page_break()

# ==================== 二、数据集介绍 ====================
add_h(doc, '二、数据集介绍', 1)

add_h(doc, '2.1 HotpotQA 数据集概述', 2)

add_p(doc, 'HotpotQA 是一个基于维基百科（Wikipedia）的大规模多跳问答数据集，由 Yang 等人在 2018 年 EMNLP 国际会议上提出。该数据集包含约 11.3 万条问答对，具有以下四个核心特征：', indent=True)

features = [
    '多跳推理需求：每个问题需要从多篇不同的维基百科文档中查找信息并进行推理才能正确回答，单篇文档不足以提供完整答案',
    '问题多样性：问题覆盖广泛的主题领域，不受限于任何预定义的知识库或知识模式，保证了数据的丰富性和挑战性',
    '句子级支持事实标注：每个问题都精确标注了支持答案的文档标题和具体句子编号，为模型训练和评估提供了可靠的监督信号',
    '两种推理类型：数据集包含桥接推理（Bridge）和比较推理（Comparison）两种类型，前者需要链式信息传递，后者需要对多个实体进行比较',
]
for feat in features:
    add_bullet(doc, '  ● ' + feat)

add_p(doc, '数据集来源：https://huggingface.co/datasets/hotpotqa/hotpot_qa', indent=True)
add_p(doc, '原始论文：Yang et al. "HotpotQA: A Dataset for Diverse, Explainable Multi-hop Question Answering", EMNLP 2018.', indent=True)

add_screenshot_placeholder(doc, '图1：HotpotQA 数据集 HuggingFace 页面截图')

add_h(doc, '2.2 数据集结构', 2)

add_p(doc, 'HotpotQA 数据集提供了两种配置（Configuration），适用于不同的研究场景：', indent=True)

add_p(doc, '（1）Distractor（干扰设置）：这是本实验选用的配置。在此设置中，每个问题除了提供支持文档外，还混入了干扰文档（Distractor Documents）。这些干扰文档与问题主题相关但不包含答案所需的关键信息，使得检索任务更具挑战性，更贴近真实世界的问答场景。', indent=True)

add_p(doc, '（2）Fullwiki（全文设置）：在此设置中，系统需要从整个维基百科（约 500 万篇文章）中检索支持文档，任务难度更大，主要用于评估检索系统的性能。', indent=True)

add_p(doc, '本实验选用 Distractor 配置的验证集（validation split），共包含 7,405 条数据。每条数据包含以下 7 个字段：')

fields = [
    ('id', '问题的唯一标识符（字符串类型）'),
    ('question', '多跳推理问题的文本内容'),
    ('answer', '问题的正确答案'),
    ('type', '推理类型，取值为 "bridge"（桥接推理）或 "comparison"（比较推理）'),
    ('level', '难度级别，取值为 "easy"、"medium" 或 "hard"'),
    ('supporting_facts', '支持事实，包含两个数组：title（支持文档标题列表）和 sent_id（支持句子编号列表）'),
    ('context', '上下文文档，包含两个数组：title（所有相关文档标题列表）和 sentences（每个文档的句子列表）'),
]
for name, desc in fields:
    add_bullet(doc, f'  ● {name}：{desc}')

add_screenshot_placeholder(doc, '图2：数据集字段结构示例截图')

doc.add_page_break()

# ==================== 三、数据获取与预处理 ====================
add_h(doc, '三、数据获取与预处理', 1)

add_h(doc, '3.1 数据获取过程', 2)

add_p(doc, '本实验通过 HuggingFace Datasets API 获取 HotpotQA 数据集。由于国内网络环境的限制，访问 HuggingFace 需要配置代理服务器。实验环境使用本地代理（127.0.0.1:7897）进行网络连接。', indent=True)

add_p(doc, '数据获取的具体步骤如下：')

steps_acq = [
    '步骤一：访问 HuggingFace API 接口（https://huggingface.co/api/datasets/hotpotqa/hotpot_qa），获取数据集的元信息，包括可用的配置（configs）、分片（splits）和数据格式',
    '步骤二：选择 "distractor" 配置和 "validation" 分片，获取 Parquet 格式数据文件的下载链接',
    '步骤三：通过 Python urllib 库，配置代理处理器（ProxyHandler），下载验证集 Parquet 文件（约 45MB）',
    '步骤四：同时下载训练集第一部分（train/0.parquet），用于了解完整的数据分布情况',
    '步骤五：使用 Python pandas 库配合 pyarrow 引擎读取 Parquet 文件，将数据加载为 DataFrame 格式',
]
for s in steps_acq:
    add_bullet(doc, s)

add_p(doc, '数据获取的核心代码实现如下：', indent=True)

add_code(doc, '''# 数据下载核心代码
import urllib.request

# 配置代理
proxy_handler = urllib.request.ProxyHandler({'https': 'http://127.0.0.1:7897'})
opener = urllib.request.build_opener(proxy_handler)

# 下载验证集 Parquet 文件
url = "https://huggingface.co/api/datasets/hotpotqa/hotpot_qa/parquet/distractor/validation/0.parquet"
resp = opener.open(url, timeout=60)
with open("data/hotpotqa_validation.parquet", "wb") as f:
    f.write(resp.read())

# 读取 Parquet 文件
import pandas as pd
df = pd.read_parquet("data/hotpotqa_validation.parquet")
print(f"数据形状: {df.shape}")  # (7405, 7)''')

add_screenshot_placeholder(doc, '图3：数据下载过程的终端输出截图')

add_h(doc, '3.2 数据预处理', 2)

add_p(doc, '下载的原始数据以 Parquet 列式存储格式保存，需要进行一系列预处理才能用于后续的 Web 检索和聚类分析。预处理的详细步骤如下：', indent=True)

prep_steps = [
    '步骤一（数据读取与格式转换）：使用 pandas 读取 Parquet 文件，将数据转换为 DataFrame 格式。由于数据中包含嵌套的 numpy 数组（如 context 字段中的 sentences 是二维数组），需要递归转换为 Python 原生列表，以便后续的 JSON 序列化。',
    '步骤二（Web 搜索数据采样）：从 7,405 条数据中抽取 500 条用于 Web 搜索功能。为保证类型的平衡性，分别抽取 bridge 类型 350 条和 comparison 类型 150 条，使两种类型的比例与原始数据基本一致。',
    '步骤三（TF-IDF 特征提取）：对所有 7,405 条问题进行文本向量化。设置最大特征数为 500、过滤英文停用词、忽略过高/过低文档频率的词汇，最终生成 7405×500 的特征矩阵。',
    '步骤四（聚类计算）：使用 K-Means 算法（k=5）对 TF-IDF 特征矩阵进行聚类，同时使用 PCA 将 500 维特征降至 2 维用于可视化展示。',
    '步骤五（数据导出）：将所有处理结果导出为 JSON 格式文件，供 Web 前端直接读取使用。',
]
for s in prep_steps:
    add_bullet(doc, s)

add_p(doc, '数据预处理的核心代码实现如下：', indent=True)

add_code(doc, '''# 数据预处理核心代码
import pandas as pd
import json

df = pd.read_parquet("data/hotpotqa_validation.parquet")

# 平衡采样：bridge 类型 350 条 + comparison 类型 150 条
bridge_df = df[df['type'] == 'bridge'].head(350)
comp_df = df[df['type'] == 'comparison'].head(150)
sample_df = pd.concat([bridge_df, comp_df])

# 导出为 JSON 格式（处理 numpy 类型）
sample_df.to_json("data/hotpotqa_data.json", orient="records", force_ascii=False)

# TF-IDF 特征提取与聚类
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
import numpy as np

questions = df['question'].tolist()
vectorizer = TfidfVectorizer(max_features=500, stop_words='english', max_df=0.8, min_df=2)
tfidf_matrix = vectorizer.fit_transform(questions)

kmeans = KMeans(n_clusters=5, random_state=42, n_init=10)
clusters = kmeans.fit_predict(tfidf_matrix.toarray())

pca = PCA(n_components=2)
coords_2d = pca.fit_transform(tfidf_matrix.toarray())''')

add_p(doc, '预处理完成后，生成了以下 5 个数据文件，供 Web 前端使用：')

data_files = [
    ('hotpotqa_data.json（3.0 MB）', '包含 500 条完整数据，每条包含问题、答案、类型、支持事实和上下文文档，用于 Web 搜索与多跳检索功能'),
    ('cluster_data.json（2.0 MB）', '包含全部 7,405 条数据的聚类结果，包括聚类标签、PCA 二维坐标等，用于聚类可视化'),
    ('cluster_terms.json（2.6 KB）', '包含每个聚类（0-4）的 Top-10 关键词及其 TF-IDF 权重，用于解释聚类含义'),
    ('questions.json（412 KB）', '包含 500 条问题摘要数据（id、question、answer、type、level），用于快速搜索匹配'),
    ('stats.json（101 B）', '包含数据集总体统计信息（总记录数、类型分布、聚类数量）'),
]
for fname, fdesc in data_files:
    add_bullet(doc, f'  ● {fname}：{fdesc}')

add_screenshot_placeholder(doc, '图4：数据预处理完成后生成的 JSON 文件列表截图')

add_screenshot_placeholder(doc, '图5：数据预处理统计信息（类型分布、聚类分布）截图')

doc.add_page_break()

# ==================== 四、多跳检索系统设计 ====================
add_h(doc, '四、多跳检索系统设计', 1)

add_h(doc, '4.1 系统架构', 2)

add_p(doc, '本系统采用纯前端架构设计，使用 HTML + CSS + JavaScript 技术栈构建单页面应用（SPA），无需任何后端服务器支持。这种架构具有部署简单、加载快速、维护方便等优点，特别适合托管在 GitHub Pages 这类静态网站托管平台上。', indent=True)

add_p(doc, '系统的整体架构包含以下四个核心模块：')

modules = [
    ('搜索与检索模块（Search & Retrieval）', '支持用户通过关键词搜索多跳问题，可根据推理类型（bridge/comparison）和聚类标签进行过滤筛选。搜索结果展示每个问题的多跳推理链路，包括支持文档列表、上下文句子以及高亮标注的支持事实。'),
    ('聚类可视化模块（Clustering Visualization）', '在 Canvas 画布上使用散点图展示 PCA 降维后的聚类结果。7405 个点代表所有问题，不同颜色代表不同聚类。支持鼠标悬停交互，显示问题的详细信息和所属聚类。'),
    ('统计分析模块（Statistics Dashboard）', '使用 Chart.js 图表库展示数据统计信息，包括问题类型分布饼图、聚类大小分布柱状图、各聚类内部类型组成的堆叠柱状图等，直观呈现数据特征。'),
    ('关于模块（About）', '介绍 HotpotQA 数据集的基本信息、系统功能说明以及技术实现细节。'),
]
for mname, mdesc in modules:
    add_bullet(doc, f'  ● {mname}：{mdesc}')

add_screenshot_placeholder(doc, '图6：Web 系统整体架构界面截图（导航栏和四个标签页）')

add_h(doc, '4.2 多跳检索功能', 2)

add_p(doc, '多跳检索是本系统最核心的功能模块，旨在直观展示多跳推理问答的完整过程。用户可以通过以下三种方式进行查询：', indent=True)

add_p(doc, '（1）关键词搜索：在搜索框中输入任意英文关键词（如 "nationality"、"film"、"director"、"born"、"football" 等），系统会在所有问题的文本内容中进行全文匹配，返回包含该关键词的问题列表。', indent=True)

add_p(doc, '（2）类型过滤：通过下拉菜单选择 "Bridge（桥接推理）" 或 "Comparison（比较推理）"，筛选出特定推理类型的问题。', indent=True)

add_p(doc, '（3）聚类过滤：通过下拉菜单选择聚类编号（0-4），筛选出属于特定语义类别的问题。', indent=True)

add_p(doc, '每条搜索结果展示以下多跳推理信息：')
retrieval_items = [
    '问题编号与问题文本：显示匹配问题的完整内容',
    '元信息标签：以彩色标签形式显示推理类型、难度级别、所属聚类编号和支持文档数量',
    '答案展示：以绿色醒目标签展示问题的正确答案',
    '支持文档链路：以箭头（→）连接的形式展示多跳推理的文档链路，清晰呈现从文档 A 到文档 B 的信息传递过程',
    '可展开的上下文详情：用户可以点击展开查看每篇文档的完整句子列表，其中被标注为支持事实的句子以黄色高亮背景突出显示，便于理解推理关键所在',
]
for item in retrieval_items:
    add_bullet(doc, f'  ● {item}')

add_p(doc, '此外，系统还实现了分页功能（每页显示 10 条结果），支持大量搜索结果的高效浏览。', indent=True)

add_screenshot_placeholder(doc, '图7：多跳检索功能 —— 搜索界面截图（输入关键词后的搜索结果列表）')
add_screenshot_placeholder(doc, '图8：多跳检索功能 —— 展开上下文详情截图（显示文档句子和支持事实高亮）')

add_h(doc, '4.3 检索示例分析', 2)

add_p(doc, '为了更好地理解多跳推理的过程，以下以一个典型的多跳问题进行详细分析：', indent=True)

add_p(doc, '示例问题："Were Scott Derrickson and Ed Wood of the same nationality?"（Scott Derrickson 和 Ed Wood 是同一国籍吗？）', indent=True)

add_p(doc, '推理类型：Comparison（比较推理）—— 需要比较两个实体的属性。', indent=True)

add_p(doc, '多跳推理过程分析：')
example_steps = [
    '第一跳 —— 检索文档 "Scott Derrickson"：系统首先查找关于 Scott Derrickson 的维基百科文档，阅读其内容，提取关键信息：Scott Derrickson 是一位美国电影导演，出生于美国科罗拉多州丹佛市。',
    '第二跳 —— 检索文档 "Ed Wood"：系统接着查找关于 Ed Wood 的维基百科文档，阅读其内容，提取关键信息：Ed Wood 是一位美国电影导演、编剧和制片人，出生于美国纽约州波基普西市。',
    '推理 —— 比较两个实体的国籍：系统从两篇文档中分别提取了两人的国籍信息，发现 Scott Derrickson 是美国人（American），Ed Wood 也是美国人（American），因此两人的国籍相同。',
    '得出结论 —— 答案："yes"：系统基于以上推理过程得出最终答案。',
]
for s in example_steps:
    add_bullet(doc, f'  ● {s}')

add_p(doc, '支持事实标注：在数据集中，两篇文档的第 0 句分别被标注为支持事实（sent_id = 0），这两句话包含了回答该问题所需的关键国籍信息。这正是 HotpotQA 数据集的重要特征——精确到句子级别的标注使得推理过程可追溯、可解释。', indent=True)

doc.add_page_break()

# ==================== 五、聚类分析方法 ====================
add_h(doc, '五、聚类分析方法', 1)

add_h(doc, '5.1 TF-IDF 特征提取', 2)

add_p(doc, 'TF-IDF（Term Frequency-Inverse Document Frequency，词频-逆文档频率）是信息检索和文本挖掘领域最常用的文本特征加权技术之一。它能够评估一个词语对于一篇文档的重要程度，其核心思想是：一个词在一篇文档中出现次数越多（TF 值高），同时在所有文档中出现次数越少（IDF 值高），则该词对该文档的区分能力越强。', indent=True)

add_p(doc, 'TF-IDF 的计算公式如下：')
add_p(doc, '  TF(t, d) = (词 t 在文档 d 中出现的次数) / (文档 d 的总词数)', indent=True)
add_p(doc, '  IDF(t) = log(总文档数 / 包含词 t 的文档数)', indent=True)
add_p(doc, '  TF-IDF(t, d) = TF(t, d) × IDF(t)', indent=True)

add_p(doc, '本实验使用 scikit-learn 库中的 TfidfVectorizer 类对 7,405 条问题进行文本特征提取。具体参数设置如下：')

tfidf_params = [
    'max_features=500：只保留 TF-IDF 权重最高的 500 个词语作为特征维度，降低计算复杂度的同时保留最重要的语义信息',
    'stop_words="english"：使用内置的英文停用词表，过滤掉 "the"、"is"、"a"、"of" 等无实际意义的常见功能词',
    'max_df=0.8：忽略在超过 80% 文档中出现的词语（高频通用词），这些词缺乏区分能力',
    'min_df=2：忽略仅在少于 2 个文档中出现的词语（极低频稀有词），这些词可能是拼写错误或噪音',
]
for p in tfidf_params:
    add_bullet(doc, f'  ● {p}')

add_p(doc, '经过 TF-IDF 特征提取后，得到 7405×500 的特征矩阵，每一行代表一个问题在 500 维语义空间中的向量表示。', indent=True)

add_code(doc, '''# TF-IDF 特征提取代码
from sklearn.feature_extraction.text import TfidfVectorizer

vectorizer = TfidfVectorizer(
    max_features=500,      # 保留 500 个最重要的特征词
    stop_words='english',   # 过滤英文停用词
    max_df=0.8,             # 忽略文档频率超过 80% 的词
    min_df=2                # 忽略文档频率低于 2 的词
)
tfidf_matrix = vectorizer.fit_transform(questions)  # 7405 × 500 矩阵''')

add_h(doc, '5.2 K-Means 聚类', 2)

add_p(doc, 'K-Means（K-均值）算法是最经典、应用最广泛的基于距离的聚类算法之一。其核心思想是将 n 个数据点划分到 k 个聚类中，使得每个数据点属于离它最近的聚类中心（均值）所代表的聚类，以最小化簇内平方误差和（Within-Cluster Sum of Squares, WCSS）。', indent=True)

add_p(doc, 'K-Means 算法的基本流程如下：')
kmeans_steps = [
    '（1）初始化：随机选择 k 个数据点作为初始聚类中心（本实验设置 k=5，使用 k-means++ 智能初始化方法）',
    '（2）分配步骤：计算每个数据点到各聚类中心的欧氏距离，将其分配到距离最近的聚类',
    '（3）更新步骤：重新计算每个聚类的中心点（该聚类所有数据点的均值）',
    '（4）迭代：重复步骤（2）和（3），直到聚类中心不再发生显著变化或达到最大迭代次数',
]
for s in kmeans_steps:
    add_bullet(doc, s)

add_p(doc, '本实验使用 scikit-learn 的 KMeans 类，设置 n_clusters=5（将问题分为 5 个语义类别），random_state=42（保证结果可复现），n_init=10（使用 10 次不同的初始化运行，选择最优结果）。', indent=True)

add_p(doc, '聚类结果及各聚类的关键词如下：')
cluster_results = [
    ('聚类 0（733 条，占比 9.9%）', '关键词：did, city, play, located, population —— 主要涉及地理位置、城市信息、人口统计等相关问题'),
    ('聚类 1（406 条，占比 5.5%）', '关键词：born, year, actor, singer, american —— 主要涉及人物出生年份、歌手演员等娱乐人物信息'),
    ('聚类 2（551 条，占比 7.4%）', '关键词：film, directed, american, actor, director —— 主要涉及电影作品、导演、演员等影视行业相关问题'),
    ('聚类 3（720 条，占比 9.7%）', '关键词：team, played, county, football, nationality —— 主要涉及运动队、国籍、足球比赛等体育领域问题'),
    ('聚类 4（4995 条，占比 67.5%）', '关键词：american, known, does, located, company —— 最大聚类，涵盖各类综合性话题'),
]
for cname, cdesc in cluster_results:
    add_bullet(doc, f'  ● {cname}：{cdesc}')

add_code(doc, '''# K-Means 聚类代码
from sklearn.cluster import KMeans

kmeans = KMeans(n_clusters=5, random_state=42, n_init=10)
clusters = kmeans.fit_predict(tfidf_matrix.toarray())
print(f"聚类分布: {np.bincount(clusters)}")  # [733, 406, 551, 720, 4995]''')

add_screenshot_placeholder(doc, '图9：聚类分析结果 —— 各聚类关键词与样本问题截图')
add_screenshot_placeholder(doc, '图10：聚类大小分布柱状图截图')

add_h(doc, '5.3 PCA 降维可视化', 2)

add_p(doc, '由于 TF-IDF 特征矩阵为 500 维高维空间，无法直接在二维平面上进行可视化。PCA（Principal Component Analysis，主成分分析）是最常用的线性降维方法，它通过正交变换将原始高维数据投影到方差最大的几个方向（主成分）上，在尽可能保留数据原始结构信息的同时降低维度。', indent=True)

add_p(doc, '本实验使用 PCA 将 500 维的 TF-IDF 特征降维到 2 维空间，然后将每个问题的二维坐标绘制在 Canvas 交互式散点图上。散点图中：')
pca_items = [
    '每个散点代表一个问题，共 7,405 个点',
    '散点的颜色表示该问题所属的聚类（5 种颜色对应 5 个聚类）',
    '鼠标悬停时显示该问题的详细信息：聚类编号、问题文本（前 120 字符）和答案',
    '散点图下方展示每个聚类的图例和样本数量',
]
for item in pca_items:
    add_bullet(doc, f'  ● {item}')

add_code(doc, '''# PCA 降维代码
from sklearn.decomposition import PCA

pca = PCA(n_components=2)
coords_2d = pca.fit_transform(tfidf_matrix.toarray())
# coords_2d 形状: (7405, 2)，可直接用于二维散点图绘制''')

add_screenshot_placeholder(doc, '图11：PCA 降维聚类散点图截图（Canvas 交互式可视化）')

doc.add_page_break()

# ==================== 六、可视化设计 ====================
add_h(doc, '六、可视化设计', 1)

add_p(doc, '数据可视化是本系统的重要组成部分，旨在帮助用户直观理解数据集的特征和聚类分析的结果。系统共设计了 4 种可视化组件，分别满足不同的分析需求：', indent=True)

vis_items = [
    ('交互式散点图（Canvas API 实现）', '这是系统的核心可视化组件。在聚类分析标签页中，使用 HTML5 Canvas 元素绘制包含 7,405 个数据点的彩色散点图。每个点的颜色对应其所属聚类，鼠标悬停时通过浮层（Tooltip）显示该问题的详细信息。散点图支持响应式布局，自适应不同屏幕尺寸。'),
    ('问题类型分布饼图（Chart.js 实现）', '使用环形图（Doughnut Chart）展示 Bridge 和 Comparison 两种推理类型的数据占比。紫色和蓝紫色配色方案与系统整体设计风格保持一致。该图直观呈现了桥接推理类型（79.9%）远多于比较推理类型（20.1%）的数据特征。'),
    ('聚类大小分布柱状图（Chart.js 实现）', '使用彩色柱状图展示 5 个聚类各自包含的问题数量。柱体的颜色与散点图中各聚类的颜色一致（红、蓝、黄、青、紫），便于跨图表对比。从该图可以直观看出聚类 4（紫色）占据了绝大多数数据（67.5%）。'),
    ('聚类-类型组成堆叠柱状图（Chart.js 实现）', '使用堆叠柱状图展示每个聚类内部 Bridge 和 Comparison 两种类型的组成比例。紫色代表 Bridge 类型，深紫色代表 Comparison 类型，使得用户能够分析各聚类的类型分布特征。'),
]
for vname, vdesc in vis_items:
    add_bullet(doc, f'  ● {vname}：{vdesc}')

add_p(doc, '所有图表均使用 Chart.js 4.4.0 库（通过 CDN 加载）实现。Chart.js 是一个轻量级、响应式的开源图表库，支持 8 种图表类型，具有良好的浏览器兼容性和动画效果。散点图则使用原生 Canvas API 实现，以获得更高的渲染性能和自定义交互能力。', indent=True)

add_screenshot_placeholder(doc, '图12：统计图表页面截图 —— 类型分布环形图和聚类分布柱状图')
add_screenshot_placeholder(doc, '图13：统计图表页面截图 —— 各聚类类型组成堆叠柱状图')

doc.add_page_break()

# ==================== 七、Web 应用部署 ====================
add_h(doc, '七、Web 应用部署', 1)

add_p(doc, '本系统选择 GitHub Pages 作为部署平台。GitHub Pages 是 GitHub 提供的免费静态网站托管服务，具有部署简单、全球 CDN 加速、HTTPS 加密等优势，非常适合托管纯前端项目。', indent=True)

add_p(doc, '部署的具体步骤如下：')

deploy_steps = [
    ('步骤一：初始化 Git 仓库', '在项目根目录执行 git init，配置用户信息（git config user.name 和 user.email），创建 .gitignore 文件排除不需要提交的临时文件（如 __pycache__、*.pyc、*.parquet 等）。'),
    ('步骤二：提交代码', '使用 git add 将 index.html、data/ 目录下的所有 JSON 数据文件、.gitignore 等文件添加到暂存区，然后使用 git commit 创建首次提交，提交信息描述项目的主要内容。'),
    ('步骤三：安装并配置 GitHub CLI', '使用 winget 包管理器安装 GitHub CLI（gh 命令），通过设备认证流程（Device Flow）完成 GitHub 账号登录。本实验使用的是账号 natrooo。'),
    ('步骤四：创建远程仓库', '使用 gh repo create 命令创建名为 "hotpotqa-analysis" 的公开仓库，同时将本地代码推送到远程（--push 参数）。配置 Git 代理（git config http.proxy）解决网络连接问题。'),
    ('步骤五：启用 GitHub Pages', '通过 GitHub API 调用，设置仓库的 Pages 源为 master 分支的根目录（"/"路径）。系统自动构建并部署，几分钟后即可通过 https://natrooo.github.io/hotpotqa-analysis/ 访问在线系统。'),
]
for sname, sdesc in deploy_steps:
    add_bullet(doc, f'  ● {sname}：{sdesc}')

add_p(doc, '部署完成后，系统可通过以下链接访问：')
add_p(doc, '  在线地址：https://natrooo.github.io/hotpotqa-analysis/', indent=True)
add_p(doc, '  GitHub 仓库：https://github.com/natrooo/hotpotqa-analysis', indent=True)

add_screenshot_placeholder(doc, '图14：GitHub 仓库页面截图（代码文件列表）')
add_screenshot_placeholder(doc, '图15：GitHub Pages 设置页面截图（显示部署成功信息）')
add_screenshot_placeholder(doc, '图16：在线系统首页截图（展示完整功能界面）')

doc.add_page_break()

# ==================== 八、实验结果与分析 ====================
add_h(doc, '八、实验结果与分析', 1)

add_h(doc, '8.1 聚类效果分析', 2)

add_p(doc, '通过 TF-IDF 特征提取和 K-Means 聚类（k=5），7,405 条多跳推理问题被成功分为 5 个语义类别。从各聚类的关键词和样本问题来看，聚类结果具有较好的语义一致性和可解释性：', indent=True)

analysis_items = [
    '聚类 0（733 条，9.9%）以 "city"、"located"、"population" 等地理相关词汇为特征，主要包含关于地理位置、行政区划、人口统计的问题。这些问题通常需要查找某个实体所在的城市、国家或地区信息。',
    '聚类 1（406 条，5.5%）以 "born"、"year"、"actor"、"singer" 等人物相关词汇为特征，主要包含关于人物出生信息的问题。聚类规模最小，说明以人物出生信息为核心的多跳问题相对较少。',
    '聚类 2（551 条，7.4%）以 "film"、"directed"、"actor"、"director" 等影视行业词汇为特征，主要包含关于电影作品和影视从业者的问题。这类问题通常需要多跳检索电影的制作信息和演职人员信息。',
    '聚类 3（720 条，9.7%）以 "team"、"played"、"football"、"nationality" 等体育相关词汇为特征，主要包含关于运动队、运动赛事和运动员国籍的问题。国籍比较类问题在此聚类中较为集中。',
    '聚类 4（4995 条，67.5%）是最大的聚类，以 "american"、"known"、"does"、"located"、"company" 等高频通用词汇为特征。该聚类的问题主题非常多样化，涵盖了从公司信息到科学概念再到历史事件的各类话题。其规模巨大（占总量的 2/3），说明多数多跳问题具有高度的主题多样性。',
]
for item in analysis_items:
    add_bullet(doc, f'  ● {item}')

add_p(doc, '聚类结果表明，TF-IDF + K-Means 的组合能够在一定程度上捕捉多跳问题的语义特征，将相似主题的问题归为一类。但也存在局限性：聚类 4 的规模过大（67.5%），说明使用 TF-IDF 词袋模型和 K-Means 线性聚类在细粒度语义区分方面能力有限。未来可以尝试使用预训练语言模型（如 BERT、Sentence-BERT）的句向量进行聚类，以获得更精细的语义表示。', indent=True)

add_h(doc, '8.2 多跳检索效果分析', 2)

add_p(doc, '本系统的多跳检索功能支持三种查询方式的组合使用，具有良好的实用性和灵活性：', indent=True)

retrieval_analysis = [
    '关键词检索效率：系统在前端使用 JavaScript 的原生字符串匹配（toLowerCase + includes），对于 500 条搜索数据，响应时间为毫秒级别，用户体验流畅。',
    '过滤准确性：通过类型过滤（bridge/comparison）和聚类过滤（0-4），用户可以精确缩小搜索范围，快速定位感兴趣的问题类别。',
    '多跳上下文展示：每条搜索结果完整展示了从文档 A 到文档 B 的推理链路，支持事实以黄色高亮标注，使得多跳推理过程一目了然。这是本系统区别于普通搜索工具的核心特色。',
    '交互体验：展开/折叠式上下文设计使得界面简洁不拥挤，用户可按需查看详细文档内容。分页功能（每页 10 条）保证了大数据量时的浏览体验。',
]
for item in retrieval_analysis:
    add_bullet(doc, f'  ● {item}')

add_screenshot_placeholder(doc, '图17：多跳检索结果展示截图（包含完整的推理链路和支持事实高亮）')

add_h(doc, '8.3 数据集统计分析', 2)

add_p(doc, '对 HotpotQA Distractor 验证集的统计分析揭示了以下数据特征：', indent=True)

stats_items = [
    '数据总量：验证集共包含 7,405 条多跳推理问题，规模适中，适合进行各类分析和实验。',
    '推理类型分布：Bridge（桥接推理）类型有 5,918 条，占比 79.9%；Comparison（比较推理）类型有 1,487 条，占比 20.1%。Bridge 类型占比接近 80%，说明大多数多跳问题需要链式信息传递（如：A 的属性确定 B，B 的属性确定答案），而非简单的实体比较。这也反映了真实世界中复杂推理问题的典型特征。',
    '难度级别：验证集中所有问题均为 "hard"（困难）级别。这与 HotpotQA 数据集的设计理念一致——验证集和测试集由更具挑战性的问题组成，用于严格评估模型的推理能力。',
    '支持文档数量：每个问题通常需要 2 篇支持文档（因为这是 2-hop 推理任务），但 context 中提供了 10 篇干扰文档，模拟了真实检索场景中的噪音环境。',
]
for item in stats_items:
    add_bullet(doc, f'  ● {item}')

add_screenshot_placeholder(doc, '图18：统计分析页面完整截图（含统计卡片和所有图表）')

doc.add_page_break()

# ==================== 九、实验总结 ====================
add_h(doc, '九、实验总结', 1)

add_p(doc, '本实验围绕多跳推理问答这一前沿研究方向，基于 HotpotQA 数据集设计并实现了一个功能完整的多跳检索与聚类分析 Web 系统。实验涵盖了数据处理、特征工程、机器学习、数据可视化和 Web 开发等多个技术领域，是一次综合性的工程实践。', indent=True)

add_p(doc, '本实验主要完成的工作包括：')

summary_items = [
    '成功获取并预处理了 HotpotQA 数据集。通过 HuggingFace API 下载了 Distractor 配置的验证集（7,405 条多跳问题），完成了 Parquet 格式解析、numpy 类型转换、平衡采样和 JSON 序列化等预处理工作。',
    '实现了 TF-IDF 文本特征提取。使用 scikit-learn 的 TfidfVectorizer，设置 500 个特征词，对全部 7,405 条问题进行了文本向量化，生成了 7405×500 维的特征矩阵。',
    '使用 K-Means 算法对问题进行了聚类分析。设置 k=5，将问题分为 5 个语义类别，并通过关键词分析验证了聚类结果的可解释性。',
    '通过 PCA 降维将高维特征降至 2 维空间，实现了聚类结果的交互式可视化展示。用户可以在散点图中直观观察问题的分布模式和聚类边界。',
    '构建了一个功能完整的 Web 交互系统。包含多跳检索、聚类可视化、统计分析和数据集介绍等 4 个核心模块，支持关键词搜索、类型过滤、聚类过滤等多种查询方式。',
    '实现了多种数据可视化图表。使用 Chart.js 库绘制了类型分布饼图、聚类分布柱状图和聚类组成堆叠柱状图，直观呈现了数据的统计特征。',
    '将系统成功部署到 GitHub Pages，实现了在线访问。通过 GitHub CLI 创建仓库、推送代码，并通过 API 启用了 Pages 服务。',
]
for item in summary_items:
    add_bullet(doc, f'  ● {item}')

add_p(doc, '通过本次实验，我获得了以下收获和体会：')

learnings = [
    '深入理解了多跳推理问答任务的特点和挑战。与单跳问答相比，多跳推理更接近人类的复杂推理过程，需要系统具备跨文档信息整合和逻辑推理的能力。HotpotQA 数据集通过精确的句子级标注为这一研究方向提供了宝贵的资源。',
    '掌握了 TF-IDF 文本特征提取和 K-Means 聚类的原理和实践。TF-IDF 虽然是一种经典的词袋模型，无法捕捉词语间的语义关系，但其计算高效、结果可解释的优点使其在文本分析中仍然具有重要的应用价值。',
    '认识到降维技术在高维数据可视化中的重要作用。PCA 通过线性变换将 500 维特征降至 2 维，虽然损失了部分信息，但使得人类能够直观地观察和理解数据的分布模式和聚类结构。',
    '提升了全栈 Web 开发能力。从数据预处理到前端界面设计，再到云端部署，完整体验了一个数据驱动 Web 应用的完整开发流程。',
    '体会到数据驱动决策的重要性。通过统计分析和可视化，能够从不同角度理解数据特征，为后续的模型选择和算法优化提供依据。',
]
for item in learnings:
    add_bullet(doc, f'  ● {item}')

add_p(doc, '未来改进方向：')
future_items = [
    '引入预训练语言模型（如 Sentence-BERT）替代 TF-IDF 进行文本语义表示，提升聚类的语义准确度',
    '增加更多交互功能，如用户自定义聚类数量、动态调整 PCA 参数等',
    '集成真实的问答模型（如基于 BERT 的 Reader），实现在线推理演示',
    '优化移动端适配，提升在小屏设备上的用户体验',
    '增加 2WikiMultihopQA 等其他多跳数据集，扩展系统的覆盖范围',
]
for item in future_items:
    add_bullet(doc, f'  ● {item}')

doc.add_page_break()

# ==================== 附录：截图清单 ====================
add_h(doc, '附录：截图清单', 1)

add_p(doc, '请按照以下清单依次添加实验过程的截图。每张截图下方请添加简短的图注说明。所有截图添加完毕后，请删除本附录中的文字说明。', indent=True)

add_p(doc, '')
add_p(doc, '第一部分：数据集相关截图')

screenshots = [
    ('图1', 'HotpotQA 数据集 HuggingFace 页面', 'https://huggingface.co/datasets/hotpotqa/hotpot_qa 页面截图，展示数据集的基本信息和配置选项'),
    ('图2', '数据集字段结构', 'Python 终端输出截图，展示数据集的列名、形状和前几条数据的字段内容'),
    ('图3', '数据下载过程', 'Python 终端输出截图，展示通过 HuggingFace API 下载 Parquet 文件的过程'),
    ('图4', '预处理生成文件', '文件资源管理器截图，展示 data/ 目录下的 5 个 JSON 文件（hotpotqa_data.json, cluster_data.json, cluster_terms.json, questions.json, stats.json）'),
    ('图5', '数据统计信息', 'Python 终端输出截图，展示数据集的类型分布（bridge: 5918, comparison: 1487）和难度级别分布'),
]
for fig_id, fig_title, fig_desc in screenshots:
    add_p(doc, f'{fig_id}：《{fig_title}》')
    add_p(doc, f'    {fig_desc}', size=11)
    add_screenshot_placeholder(doc, f'请在此处插入{fig_id}')
    add_p(doc, '')

add_p(doc, '第二部分：Web 系统功能截图')
screenshots2 = [
    ('图6', '系统整体界面', 'Web 系统首页截图，展示导航栏的四个标签页（搜索与检索、聚类可视化、统计分析、关于数据集）和顶部统计卡片'),
    ('图7', '多跳检索——搜索结果', '输入关键词 "nationality" 后的搜索结果列表截图，展示问题、类型标签、答案和推理链路'),
    ('图8', '多跳检索——上下文展开', '点击"查看多跳上下文"后展开的详细内容截图，展示文档句子列表和支持事实黄色高亮标注'),
    ('图9', '聚类关键词展示', '聚类可视化页面的关键词区域截图，展示 5 个聚类的 Top-10 关键词及其 TF-IDF 权重'),
    ('图10', '聚类大小分布图表', '统计页面截图，展示各聚类问题数量的彩色柱状图'),
    ('图11', 'PCA 聚类散点图', '聚类可视化页面的 Canvas 散点图截图，展示 7405 个彩色散点和鼠标悬停时的 tooltip 信息'),
    ('图12', '类型分布与聚类分布图表', '统计页面截图，同时展示问题类型分布环形图和聚类大小柱状图'),
    ('图13', '聚类-类型组成图表', '统计页面截图，展示各聚类内部 Bridge/Comparison 类型的堆叠柱状图'),
]
for fig_id, fig_title, fig_desc in screenshots2:
    add_p(doc, f'{fig_id}：《{fig_title}》')
    add_p(doc, f'    {fig_desc}', size=11)
    add_screenshot_placeholder(doc, f'请在此处插入{fig_id}')
    add_p(doc, '')

add_p(doc, '第三部分：部署相关截图')
screenshots3 = [
    ('图14', 'GitHub 仓库页面', 'https://github.com/natrooo/hotpotqa-analysis 页面截图，展示代码文件列表（index.html, data/ 目录等）'),
    ('图15', 'GitHub Pages 设置', '仓库 Settings > Pages 页面的截图，展示 GitHub Pages 已启用且源设置为 master 分支'),
    ('图16', '在线系统首页', 'https://natrooo.github.io/hotpotqa-analysis/ 在线访问截图，展示完整功能界面'),
]
for fig_id, fig_title, fig_desc in screenshots3:
    add_p(doc, f'{fig_id}：《{fig_title}》')
    add_p(doc, f'    {fig_desc}', size=11)
    add_screenshot_placeholder(doc, f'请在此处插入{fig_id}')
    add_p(doc, '')

# ==================== 保存 ====================
output_path = "C:/Users/LTZ/Desktop/数据2302葛超瑜实验报告.docx"
doc.save(output_path)
print(f"报告已保存至：{output_path}")
print(f"文件大小：{os.path.getsize(output_path) / 1024:.1f} KB")
print(f"总段落数：{len(doc.paragraphs)}")
