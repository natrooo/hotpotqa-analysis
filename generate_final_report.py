# -*- coding: utf-8 -*-
"""生成模仿 NewSQL 报告格式的实验报告"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ---- 页面设置 ----
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ---- 样式 ----
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)

def H(text, level=3):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
    return h

def P(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    for line in text.strip().split('\n'):
        if p.runs:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def screenshot(label):
    """插入截图占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（' + label + '）')
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============================================================
# 标题
# ============================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SQLite数据库部署与多跳推理Web应用实验报告')
run.font.name = '黑体'
run.font.size = Pt(22)
run.bold = True

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('数据2302班  葛超瑜  2026年6月8日')
run.font.name = '宋体'
run.font.size = Pt(14)

doc.add_paragraph()

# ============================================================
# 一、实验内容
# ============================================================
H('一、实验内容', 3)

P('本次实验为本课程数据库系列实验的综合应用项目。实验选取 HotpotQA 多跳推理问答数据集（约11.3万条基于维基百科的多跳问答对），在本地 Windows 环境下搭建 SQLite 关系型数据库，将7,405条多跳问答数据导入数据库中，构建包含问题表（questions）、文档上下文表（contexts）、句子表（sentences）、聚类结果表（clusters）以及全文搜索索引（FTS5）的完整数据库模式。随后使用 Python Flask 框架开发 Web 应用，通过 SQL 查询实现多跳检索、全文搜索、聚类可视化和统计分析功能，最终完成一个集数据存储、查询检索、机器学习聚类和可视化展示于一体的完整系统。')

P('实验的核心内容包括：（1）HotpotQA 数据集的获取与预处理；（2）SQLite 数据库模式设计与建表；（3）大规模数据批量导入与验证；（4）TF-IDF 文本特征提取与 K-Means 聚类计算；（5）PCA 降维与聚类结果存储；（6）Flask Web 应用开发，实现多跳检索 API、统计 API、聚类数据 API 和 SQL 查询接口；（7）前端可视化界面开发，包括散点图、饼图、柱状图和堆叠柱状图。')

# ============================================================
# 二、实验环境
# ============================================================
H('二、实验环境', 3)

env_items = [
    '操作系统：Windows 11 Home China（版本 10.0.26200）',
    '数据库：SQLite 3.35.5（内置于 Python 标准库，零配置部署）',
    '数据库文件：hotpotqa.db（50.9 MB，包含 7,405 条问题、73,700 条文档上下文、306,487 条句子）',
    'Python 环境：Python 3.9.5，主要依赖包：pandas 2.3.3、scikit-learn 1.6.1、Flask 3.1.3',
    'Web 框架：Flask + 原生 HTML/CSS/JavaScript',
    '可视化库：Chart.js 4.4.0（CDN加载）',
    '数据集：HotpotQA Distractor 验证集（7,405条多跳问答对，来自 HuggingFace）',
    '数据格式：原始 Parquet → pandas DataFrame → SQLite 关系表',
]
for item in env_items:
    P('● ' + item)

# ============================================================
# 三、实验过程与步骤
# ============================================================
H('三、实验过程与步骤', 3)

# --- (一) ---
H('（一）HotpotQA 数据集获取与解析', 4)

P('首先通过代理服务器访问 HuggingFace Datasets API，获取 HotpotQA 数据集的元信息和文件下载链接。HotpotQA 数据集提供两种配置：Distractor（干扰设置）和 Fullwiki（全文设置）。本实验选用 Distractor 配置的验证集分片（validation split），共 7,405 条数据。数据以 Parquet 列式存储格式提供，使用 Python 的 pandas 库配合 pyarrow 引擎进行读取。')

P('数据获取的核心代码如下：')

code('''import urllib.request
import pandas as pd

# 通过代理访问 HuggingFace API
proxy_handler = urllib.request.ProxyHandler({'https': 'http://127.0.0.1:7897'})
opener = urllib.request.build_opener(proxy_handler)

# 下载验证集 Parquet 文件
url = "https://huggingface.co/api/datasets/hotpotqa/hotpot_qa/parquet/distractor/validation/0.parquet"
resp = opener.open(url, timeout=60)
with open("data/hotpotqa_validation.parquet", "wb") as f:
    f.write(resp.read())

# 读取数据
df = pd.read_parquet("data/hotpotqa_validation.parquet")
print(df.shape)  # (7405, 7)
print(df.columns.tolist())  # ['id','question','answer','type','level','supporting_facts','context']''')

screenshot('图1：HuggingFace 数据集页面截图')

P('数据集包含7个字段：id（问题唯一标识）、question（多跳推理问题文本）、answer（答案）、type（推理类型：bridge/comparison）、level（难度级别）、supporting_facts（支持事实，包含文档标题和句子编号）、context（上下文文档，每个问题附带10篇文档的全部句子）。')

screenshot('图2：Python终端输出——数据形状、列名和前两条数据预览')

# --- (二) ---
H('（二）SQLite 数据库模式设计与建表', 4)

P('针对 HotpotQA 数据的多层次嵌套结构，设计了以下关系型数据库模式。将原始的一对多嵌套结构拆分为三张核心表，外加聚类结果表和全文搜索虚拟表，共五张表：')

P('（1）questions 表（问题主表）：以 id 为主键，存储 question（问题文本）、answer（答案）、type（推理类型，使用 CHECK 约束限制取值）、level（难度级别）。这是数据查询的主入口。')

P('（2）contexts 表（文档上下文表）：以自增 id 为主键，通过 question_id 外键关联到 questions 表。每条记录代表一个问题下的一个文档上下文，包含 doc_index（文档序号）、doc_title（文档标题）和 is_supporting（是否为支持文档的标志位）。')

P('（3）sentences 表（句子表）：以自增 id 为主键，通过 context_id 外键关联到 contexts 表。每条记录代表一篇文档中的一个句子，包含 sent_index（句子序号）、sent_text（句子文本）和 is_supporting_fact（是否为支持事实的标志位）。')

P('（4）clusters 表（聚类结果表）：以 question_id 为主键关联到 questions 表，存储 cluster_id（聚类编号 0-4）、coord_x 和 coord_y（PCA 降维后的二维坐标），用于前端聚类可视化。')

P('（5）questions_fts 表（全文搜索虚拟表）：使用 SQLite FTS5 引擎，对 question 和 answer 字段建立全文索引，支持高效的模糊搜索。')

P('建表 SQL 语句如下：')

code('''CREATE TABLE questions (
    id TEXT PRIMARY KEY,
    question TEXT NOT NULL,
    answer TEXT NOT NULL,
    type TEXT NOT NULL CHECK(type IN ('bridge', 'comparison')),
    level TEXT NOT NULL
);

CREATE TABLE contexts (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    question_id TEXT NOT NULL,
    doc_index INTEGER NOT NULL,
    doc_title TEXT NOT NULL,
    is_supporting INTEGER DEFAULT 0,
    FOREIGN KEY (question_id) REFERENCES questions(id)
);

CREATE TABLE sentences (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    context_id INTEGER NOT NULL,
    sent_index INTEGER NOT NULL,
    sent_text TEXT NOT NULL,
    is_supporting_fact INTEGER DEFAULT 0,
    FOREIGN KEY (context_id) REFERENCES contexts(id)
);

CREATE TABLE clusters (
    question_id TEXT PRIMARY KEY,
    cluster_id INTEGER NOT NULL,
    coord_x REAL NOT NULL,
    coord_y REAL NOT NULL,
    FOREIGN KEY (question_id) REFERENCES questions(id)
);

CREATE VIRTUAL TABLE questions_fts USING fts5(
    question, answer, content='questions', content_rowid='rowid'
);''')

screenshot('图3：SQLite 建表执行的终端截图')

# --- (三) ---
H('（三）数据批量导入与验证', 4)

P('数据导入是整个实验中最关键的步骤之一。将 7,405 条 HotpotQA 数据从 Parquet 格式导入到 SQLite 关系型数据库，需要遍历原始数据的嵌套结构，将每个问题的 context 数组（包含 10 篇文档）和每篇文档的 sentences 数组（每篇约 4-5 句）逐条展开并插入对应的表中。导入过程采用每 2,000 条提交一次事务的策略，平衡了内存占用和写入效率。')

P('导入的核心逻辑如下：')

code('''import sqlite3, pandas as pd

conn = sqlite3.connect("hotpotqa.db")
df = pd.read_parquet("data/hotpotqa_validation.parquet")

for i in range(len(df)):
    row = df.iloc[i]
    qid = str(row['id'])

    # 插入问题主表
    cursor.execute(
        "INSERT INTO questions (id, question, answer, type, level) VALUES (?,?,?,?,?)",
        (qid, str(row['question']), str(row['answer']), str(row['type']), str(row['level']))
    )

    # 遍历上下文文档
    ctx = row['context']
    sf_titles = row['supporting_facts']['title'].tolist()
    sf_sent_ids = row['supporting_facts']['sent_id'].tolist()

    for doc_idx in range(len(ctx['title'])):
        doc_title = str(ctx['title'][doc_idx])
        is_supporting = 1 if doc_title in sf_titles else 0

        cursor.execute(
            "INSERT INTO contexts (question_id, doc_index, doc_title, is_supporting) VALUES (?,?,?,?)",
            (qid, doc_idx, doc_title, is_supporting)
        )
        context_id = cursor.lastrowid

        # 遍历该文档的每个句子
        sentences = ctx['sentences'][doc_idx]
        for sent_idx, sent_text in enumerate(sentences):
            is_sf = 0
            if doc_title in sf_titles:
                sf_pos = sf_titles.index(doc_title)
                if sf_pos < len(sf_sent_ids) and sf_sent_ids[sf_pos] == sent_idx:
                    is_sf = 1
            cursor.execute(
                "INSERT INTO sentences (context_id, sent_index, sent_text, is_supporting_fact) VALUES (?,?,?,?)",
                (context_id, sent_idx, str(sent_text), is_sf)
            )

    if (i + 1) % 2000 == 0:
        conn.commit()  # 每 2000 条提交一次

conn.commit()

# 验证数据完整性
cursor.execute("SELECT COUNT(*) FROM questions")   # 7405
cursor.execute("SELECT COUNT(*) FROM contexts")     # 73700
cursor.execute("SELECT COUNT(*) FROM sentences")    # 306487''')

P('导入完成后，通过 COUNT 查询验证数据完整性：questions 表 7,405 条（与原始数据一致）、contexts 表 73,700 条（每个问题 10 篇文档，部分问题文档数略有差异）、sentences 表 306,487 条。随后构建 FTS5 全文索引，将 questions 表的内容同步到 questions_fts 虚拟表中。')

screenshot('图4：数据导入过程终端截图（显示进度和最终统计）')
screenshot('图5：COUNT 验证查询结果截图')

# --- (四) ---
H('（四）TF-IDF 特征提取与 K-Means 聚类', 4)

P('为了对 7,405 条多跳问题进行语义聚类分析，首先需要将文本转换为数值特征向量。使用 scikit-learn 的 TfidfVectorizer 进行 TF-IDF（词频-逆文档频率）特征提取，参数设置为：max_features=500（保留500个最重要的特征词）、stop_words=\'english\'（过滤英文停用词）、max_df=0.8（忽略文档频率超过80%的通用词）、min_df=2（忽略仅在少于2篇文档中出现的稀有词）。最终生成 7405×500 维的特征矩阵。')

P('随后使用 K-Means 算法（k=5，random_state=42，n_init=10）对特征矩阵进行聚类。聚类结果分布如下：聚类0（733条，主要涉及地理位置类问题）、聚类1（406条，主要涉及人物出生信息）、聚类2（551条，主要涉及影视行业）、聚类3（720条，主要涉及体育和国籍）、聚类4（4995条，涵盖各类综合话题）。')

P('由于 500 维特征无法直接可视化，使用 PCA（主成分分析）将特征降至 2 维空间，降维后的坐标存入 clusters 表，供前端 Canvas 散点图使用。聚类结果和 PCA 坐标的计算代码如下：')

code('''from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA

vectorizer = TfidfVectorizer(max_features=500, stop_words='english', max_df=0.8, min_df=2)
tfidf_matrix = vectorizer.fit_transform(questions)   # 7405 x 500

kmeans = KMeans(n_clusters=5, random_state=42, n_init=10)
clusters = kmeans.fit_predict(tfidf_matrix.toarray())

pca = PCA(n_components=2)
coords_2d = pca.fit_transform(tfidf_matrix.toarray())

# 存入数据库
for i in range(len(ids)):
    cursor.execute(
        "INSERT INTO clusters (question_id, cluster_id, coord_x, coord_y) VALUES (?,?,?,?)",
        (ids[i], int(clusters[i]), float(coords_2d[i][0]), float(coords_2d[i][1]))
    )''')

screenshot('图6：TF-IDF + K-Means + PCA 计算过程的终端输出截图')

# --- (五) ---
H('（五）Flask Web 应用开发', 4)

P('Web 后端使用 Flask 框架开发，通过 Python 内置的 sqlite3 模块连接数据库。所有 API 返回 JSON 格式数据，供前端 JavaScript 异步加载。应用共实现了以下 6 个核心 API 路由：')

P('（1）/ 首页路由：渲染系统主页面，包含多跳检索、聚类可视化、统计分析和 SQL 查询四个功能标签页。')

P('（2）/api/search 多跳检索 API：接收关键词（q）、推理类型（type）、聚类编号（cluster）和分页参数（page），使用 SQLite FTS5 全文搜索引擎在 questions_fts 表中进行模糊匹配。查询结果包含每个问题的完整多跳上下文——通过 JOIN 三张表（questions → contexts → sentences），返回每篇文档的标题和全部句子，并标注 is_supporting 和 is_supporting_fact 标志位，使前端能够高亮显示支持文档和支持事实，直观展示多跳推理链路。')

P('（3）/api/clusters 聚类数据 API：返回全部 7,405 条数据的聚类标签和 PCA 二维坐标，以及每个聚类的 Top-10 关键词及其 TF-IDF 权重，供前端 Canvas 散点图和关键词展示使用。')

P('（4）/api/stats 统计 API：执行多条聚合查询（GROUP BY type、GROUP BY cluster_id），返回类型分布、聚类分布和各聚类内的类型组成数据，供 Chart.js 图表使用。')

P('（5）/api/sql SQL 查询接口：接收用户在前端输入的 SELECT 语句，直接在 SQLite 上执行并返回查询结果。此接口仅限于 SELECT 查询，确保数据库安全。')

P('（6）/api/count 统计接口：返回三张核心表的记录总数，用于验证数据完整性。')

P('Flask 应用的核心代码结构如下：')

code('''from flask import Flask, render_template, request, jsonify
import sqlite3

app = Flask(__name__)
DB_PATH = "hotpotqa.db"

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row   # 使查询结果支持字典式访问
    return conn

@app.route('/api/search')
def search():
    keyword = request.args.get('q', '')
    # 使用 FTS5 全文搜索引擎进行模糊匹配
    if keyword:
        cur.execute("""
            SELECT q.*, c.cluster_id FROM questions q
            LEFT JOIN clusters c ON q.id = c.question_id
            WHERE q.id IN (
                SELECT rowid FROM questions_fts WHERE questions_fts MATCH ?
            )
        """, (keyword,))
    # ... 然后 JOIN contexts 和 sentences 获取完整多跳上下文

@app.route('/api/clusters')
def get_clusters():
    cur.execute("""
        SELECT q.question, q.answer, c.cluster_id, c.coord_x, c.coord_y
        FROM clusters c JOIN questions q ON c.question_id = q.id
    """)
    return jsonify({'points': [dict(r) for r in cur.fetchall()]})''')

screenshot('图7：Flask 应用代码结构截图（app.py 核心路由）')
screenshot('图8：Flask 应用启动成功终端截图')

# ============================================================
# 四、Web 功能展示
# ============================================================
H('四、Web 功能展示', 3)

P('Flask 应用启动后，在浏览器中访问 http://127.0.0.1:5000 即可打开系统首页。首页顶部展示数据库统计卡片（总问题数 7,405、Bridge 类型 5,918、Comparison 类型 1,487、数据库大小 50.9MB），下方为四个功能标签页。')

screenshot('图9：Web 首页完整截图（多跳检索标签页，含统计卡片）')

# --- 1. 多跳检索功能 ---
H('1. 多跳检索功能', 4)

P('在多跳检索标签页，用户可以通过关键词搜索问题。例如输入关键词 "nationality" 并点击搜索，系统通过 FTS5 全文搜索引擎在所有问题文本中查找包含该词的问题，并以列表形式展示。每条搜索结果包含：问题编号与文本、推理类型标签（桥接推理/比较推理）、难度级别、所属聚类编号、支持文档数量，以及以绿色标签展示的正确答案。')

P('搜索结果的底部展示了多跳推理链路（如 "Scott Derrickson → Ed Wood"），以箭头连接支持文档标题，清晰展示从第一跳文档到第二跳文档的信息传递过程。点击"查看多跳上下文"，可展开查看每篇文档的完整句子列表，其中被标注为支持事实（supporting fact）的句子以黄色高亮背景显示。每个问题默认附带 10 篇上下文文档（包含干扰文档），其中 2 篇为支持文档（以 ⭐ 标记）。')

P('系统支持三种过滤方式的组合使用：推理类型过滤（Bridge/Comparison）、聚类过滤（聚类0-4），以及关键词搜索。搜索结果支持分页浏览，每页显示 10 条结果。')

screenshot('图10：多跳检索——输入"nationality"搜索后的结果列表截图')
screenshot('图11：多跳检索——展开上下文详情截图（黄色高亮为支持事实句子）')

# --- 2. 聚类可视化功能 ---
H('2. 聚类可视化功能', 4)

P('在聚类可视化标签页，使用 HTML5 Canvas 绘制包含 7,405 个数据点的交互式散点图。每个点代表一个问题，颜色表示其 K-Means 聚类归属（红、蓝、黄、青、紫五种颜色分别对应聚类0-4）。散点图的坐标由 PCA 将 500 维 TF-IDF 特征降至 2 维得到。鼠标悬停在任何数据点上时，弹出 tooltip 显示该问题的聚类编号、问题文本摘要和答案。')

P('散点图下方展示聚类图例和各聚类的 Top-10 关键词（含 TF-IDF 权重）。从关键词可以看出：聚类0以 "city"、"located"、"population" 等地理词汇为主；聚类1以 "born"、"year"、"actor"、"singer" 等人物出生信息为主；聚类2以 "film"、"directed"、"actor"、"director" 等影视词汇为主；聚类3以 "team"、"played"、"football"、"nationality" 等体育词汇为主；聚类4以 "american"、"known"、"does"、"company" 等通用词汇为主。')

screenshot('图12：聚类可视化散点图全貌截图')
screenshot('图13：聚类可视化——鼠标悬停显示详情截图')

# --- 3. 统计分析功能 ---
H('3. 统计分析功能', 4)

P('在统计分析标签页，使用 Chart.js 库绘制了三张统计图表：')

P('（1）问题类型分布环形图：展示 Bridge 桥接推理（5,918条，79.9%）和 Comparison 比较推理（1,487条，20.1%）的占比。Bridge 类型占比接近 80%，说明大多数多跳问题需要链式信息传递，而非简单的实体比较。')

P('（2）聚类大小分布柱状图：以五种颜色展示各聚类的问题数量。聚类4（紫色）以 4,995 条占据绝对多数（67.5%），其余四个聚类的规模相近（406-733条之间）。')

P('（3）各聚类内部推理类型组成堆叠柱状图：展示每个聚类中 Bridge 和 Comparison 两种类型的组成比例。从图中可以看出，所有聚类中 Bridge 类型都占据主导地位，但聚类1和聚类2中 Comparison 类型的比例相对较高。')

screenshot('图14：统计分析页面——类型分布环形图和聚类大小柱状图截图')
screenshot('图15：统计分析页面——各聚类内部类型组成堆叠柱状图截图')

# --- 4. SQL 查询功能 ---
H('4. SQL 查询功能', 4)

P('在 SQL 查询标签页，用户可以直接输入 SELECT 语句对 SQLite 数据库进行查询。界面预设了一条示例查询（展示问题、答案、类型和聚类编号的联合查询），点击"执行查询"按钮即可查看结果。查询结果以表格形式展示在页面下方，并显示返回的行数。为保障数据库安全，此接口仅允许执行 SELECT 语句，任何非 SELECT 操作都会被拒绝。')

screenshot('图16：SQL 查询页面——执行示例查询的结果截图')

# ============================================================
# 五、实验结果、问题及解决过程
# ============================================================
H('五、实验结果、问题及解决过程', 3)

P('问题一：HuggingFace 数据访问受限')

P('由于国内网络环境的限制，直接访问 HuggingFace 的 API 接口会遭遇连接被重置（WinError 10054）的问题。最初尝试使用系统代理（127.0.0.1:7897）时，Python 的 urllib 能正确通过代理访问 HTTPS 资源，但 pip 安装包时由于清华镜像源的 SSL 证书问题而失败。')

P('解决过程：通过分析发现问题出在两个层面——Python 脚本的网络请求和 pip 包管理器的下载渠道。对于 Python 脚本，通过配置 ProxyHandler 正确使用本地代理即可正常访问 HuggingFace API。对于 pip 安装，需要在环境变量中正确设置 HTTPS_PROXY=http://127.0.0.1:7897，并使用 PyPI 官方源（https://pypi.org/simple/）而非清华镜像源。通过这种双轨并行的代理配置策略，成功安装了 pyarrow、Flask 等依赖包，并下载了 HotpotQA 数据集。')

P('问题二：数据导入过程中的 numpy 类型序列化')

P('HotpotQA 数据集中，context 字段的 sentences 是二维 numpy 数组（ndarray 嵌套 ndarray），supporting_facts 字段的 title 和 sent_id 也是 numpy 数组。直接使用 pandas 的 to_json 方法导出时，会因为 ndarray 类型不可 JSON 序列化而报错。')

P('解决过程：编写了递归的类型转换函数，在遍历数据时将 numpy 数组逐层转换为 Python 原生列表（tolist()），将 numpy 整数和浮点数转换为 Python 原生 int 和 float。同时，在 SQLite 导入时通过 str() 函数确保所有文本字段为字符串类型，避免了类型不匹配的问题。最终数据成功导入，通过 COUNT 查询验证了数据完整性——questions 表 7,405 条、contexts 表 73,700 条、sentences 表 306,487 条，与预期完全一致。')

P('问题三：FTS5 全文搜索的中文和特殊字符支持')

P('在设计全文搜索功能时，发现 SQLite FTS5 的默认分词器对英文效果最好，而 HotpotQA 的问题文本全部为英文，因此 FTS5 的默认 tokenizer 完全适用。但对于包含特殊字符（如引号、括号）的关键词，需要使用 FTS5 的转义语法进行处理。')

P('解决过程：在前端搜索框中，系统自动处理用户输入的关键词。FTS5 的 MATCH 语法支持直接传入包含空格的短语（会自动分词），无需额外处理。通过在 API 层将关键词直接传递给 FTS5 的 MATCH 子句，实现了高效的全文搜索。相比简单的 LIKE 模糊匹配，FTS5 的搜索速度更快，且支持布尔操作符和前缀匹配。')

# ============================================================
# 六、实验总结与心得
# ============================================================
H('六、实验总结与心得', 3)

P('SQLite 虽小，却在本实验中展现了令人印象深刻的实力。作为本课程数据库系列实验的综合应用，本次实验让我将之前学到的数据库设计、SQL 查询、Python 编程和 Web 开发知识整合到了一个完整的项目中。从数据获取、数据库建模、批量导入，到 Web API 开发和前端可视化，完整地走了一遍"数据→数据库→应用→展示"的全流程。')

P('选择 SQLite 作为本次实验的数据库是一个务实的决定。与之前实验中使用的 CockroachDB 等分布式 NewSQL 数据库不同，SQLite 的最大优势在于零配置——无需安装服务器、无需配置端口、无需管理用户权限，一个文件就是整个数据库。对于本实验 50.9MB 的数据库规模（7,405 条问题、30 万+条句子），SQLite 的处理能力绰绰有余。在开发过程中，Python 内置的 sqlite3 模块让数据库操作异常简洁——不需要连接池、不需要 ORM，直接写原生 SQL 就能完成所有操作。这种轻量级的方式非常适合中小规模的数据分析和 Web 演示项目。')

P('多跳推理是本次实验的核心概念，也是让我感受最深的部分。传统数据库实验通常处理的是扁平的结构化数据（如用户表、订单表），而 HotpotQA 的数据本质上是多层次的文档结构——每个问题关联 10 篇文档，每篇文档包含多个句子，且需要标注哪些文档是支持文档、哪些句子是支持事实。将这种嵌套结构映射到关系型数据库的三张表（questions → contexts → sentences），并通过外键和 JOIN 查询来还原多跳推理过程，让我深刻体会到关系模型在表达复杂数据结构时的灵活性和表达能力。')

P('在众多学过的数据库类型中——Redis 追求速度但数据模型简单，MongoDB 灵活但对复杂查询支持有限，CockroachDB 分布式能力强但部署复杂——SQLite 在"简单"和"强大"之间找到了优雅的平衡点。它保持了完整的 SQL 标准支持（包括 JSON 函数、全文搜索、窗口函数等高级特性），同时将使用门槛降到了最低。这次实验中，FTS5 全文搜索引擎的表现尤其令人惊喜——无需额外安装 Elasticsearch 或 Whoosh，SQLite 内置的 FTS5 就能提供高效的全文搜索能力，与 SQL 查询无缝集成。')

P('通过本次实验，我不仅掌握了 SQLite 数据库的建表、导入、查询和全文搜索等核心操作，还实践了 TF-IDF 特征提取、K-Means 聚类和 PCA 降维等机器学习技术，并成功构建了一个功能完整的 Flask Web 应用。这一综合性实验让我对"数据驱动的 Web 应用"有了更全面的理解，也为今后处理更大规模的数据集和更复杂的数据库系统打下了坚实基础。')

# ============================================================
# 保存
# ============================================================
output_path = "C:/Users/LTZ/Desktop/数据2302葛超瑜实验报告_final.docx"
# 如果旧文件存在且被占用，使用新文件名
try:
    doc.save(output_path)
    print(f"报告已保存至：{output_path}")
except PermissionError:
    output_path = "C:/Users/LTZ/Desktop/数据2302葛超瑜实验报告_new.docx"
    doc.save(output_path)
    print(f"原文件被占用，已保存至：{output_path}")

size_kb = os.path.getsize(output_path) / 1024
print(f"文件大小：{size_kb:.1f} KB")
